"""
Converts RiskWaiver360_Final_Documentation.md to a styled DOCX.
Uses python-docx to build the document from parsed markdown lines.
"""
import re, pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = pathlib.Path(__file__).parent / "RiskWaiver360_Final_Documentation.md"
DOCX_PATH = pathlib.Path(__file__).parent / "RiskWaiver360_Final_Documentation.docx"

RED       = RGBColor(0xDC, 0x26, 0x26)
DARK_RED  = RGBColor(0x99, 0x1B, 0x1B)
DARK      = RGBColor(0x1E, 0x29, 0x3B)
MUTED     = RGBColor(0x47, 0x55, 0x69)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xFF, 0xF7, 0xF7)
CODE_BG   = RGBColor(0x0F, 0x17, 0x2A)
CODE_FG   = RGBColor(0xE2, 0xE8, 0xF0)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else None
    if level == 1:
        p.runs[0].font.color.rgb = DARK_RED
        p.runs[0].font.size = Pt(20)
    elif level == 2:
        if run:
            run.font.color.rgb = DARK_RED
            run.font.size = Pt(14)
    elif level == 3:
        if run:
            run.font.color.rgb = DARK
            run.font.size = Pt(12)
    return p


def add_code_block(doc: Document, lines: list[str]):
    text = "\n".join(lines)
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.font.name     = "Courier New"
    run.font.size     = Pt(8.5)
    run.font.color.rgb = CODE_FG
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "0F172A")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    # left border
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),  "single")
    left.set(qn("w:sz"),   "24")
    left.set(qn("w:space"),"4")
    left.set(qn("w:color"),"DC2626")
    pBdr.append(left)
    pPr.append(pBdr)


def parse_inline(run_adder, text: str):
    """Add a paragraph run handling **bold** and `code` inline."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|`([^`]+)`)")
    last    = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            run_adder(text[last:m.start()], bold=False, code=False)
        if m.group(2):
            run_adder(m.group(2), bold=True, code=False)
        elif m.group(3):
            run_adder(m.group(3), bold=False, code=True)
        last = m.end()
    if last < len(text):
        run_adder(text[last:], bold=False, code=False)


def add_paragraph_with_inline(doc: Document, text: str, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)

    def adder(t, bold, code):
        r = p.add_run(t)
        r.bold = bold
        if code:
            r.font.name      = "Courier New"
            r.font.size      = Pt(9)
            r.font.color.rgb = RED

    parse_inline(adder, text)
    return p


def add_table_md(doc: Document, header_row: list[str], rows: list[list[str]]):
    col_count = max(len(header_row), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        if i >= col_count:
            break
        cell = hdr_cells[i]
        cell.text = h.strip().strip("*")
        set_cell_bg(cell, "991B1B")
        for run in cell.paragraphs[0].runs:
            run.font.bold      = True
            run.font.color.rgb = WHITE
            run.font.size      = Pt(9)
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(3)

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg    = "FFF7F7" if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row):
            if ci >= col_count:
                break
            cell = cells[ci]
            cell.text = val.strip()
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()


def build_docx(md_text: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Cover page
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("RiskWaiver360")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = DARK_RED
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("GRC Process Exception & Policy Waiver Management Platform")
    r2.font.size = Pt(14); r2.bold = True; r2.font.color.rgb = DARK_RED
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run("Final Hackathon Submission Documentation")
    r3.font.size = Pt(11); r3.font.color.rgb = MUTED
    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = sub3.add_run("Société Générale Technology Hackathon — PS1")
    r4.font.size = Pt(11); r4.font.color.rgb = MUTED
    doc.add_page_break()

    lines     = md_text.splitlines()
    i         = 0
    in_code   = False
    code_buf  = []
    in_table  = False
    tbl_hdr   = []
    tbl_rows  = []

    def flush_table():
        nonlocal in_table, tbl_hdr, tbl_rows
        if tbl_hdr:
            add_table_md(doc, tbl_hdr, tbl_rows)
        in_table = False; tbl_hdr = []; tbl_rows = []

    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()

        # --- Code block ---
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []; in_code = False
            else:
                if in_table: flush_table()
                in_code = True
            i += 1; continue

        if in_code:
            code_buf.append(raw)
            i += 1; continue

        # --- Table ---
        if line.startswith("|"):
            cells = [c for c in line.split("|") if c.strip() != ""]
            # separator row
            if all(re.match(r"^[-:]+$", c.strip()) for c in cells):
                i += 1; continue
            if not in_table:
                in_table = True; tbl_hdr = cells; tbl_rows = []
            else:
                tbl_rows.append(cells)
            i += 1; continue
        else:
            if in_table: flush_table()

        # Skip horizontal rule
        if re.match(r"^-{3,}$", line) or re.match(r"^\*{3,}$", line):
            i += 1; continue

        # --- Headings ---
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            add_heading(doc, text, level)
            i += 1; continue

        # --- Blank line ---
        if not line:
            i += 1; continue

        # --- Bullet list ---
        m2 = re.match(r"^[-*]\s+(.*)", line)
        if m2:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)

            def adder_li(t, bold, code, _p=p):
                r = _p.add_run(t)
                r.bold = bold
                if code:
                    r.font.name = "Courier New"; r.font.size = Pt(9); r.font.color.rgb = RED

            parse_inline(adder_li, m2.group(1))
            i += 1; continue

        # --- Numbered list ---
        m3 = re.match(r"^\d+\.\s+(.*)", line)
        if m3:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)

            def adder_ol(t, bold, code, _p=p):
                r = _p.add_run(t)
                r.bold = bold
                if code:
                    r.font.name = "Courier New"; r.font.size = Pt(9); r.font.color.rgb = RED

            parse_inline(adder_ol, m3.group(1))
            i += 1; continue

        # --- Blockquote ---
        if line.startswith(">"):
            content = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(content)
            r.italic = True; r.font.color.rgb = MUTED
            i += 1; continue

        # --- Normal paragraph ---
        add_paragraph_with_inline(doc, line)
        i += 1

    if in_table:
        flush_table()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    return doc


def main():
    print(f"Input  : {MD_PATH}")
    print(f"Output : {DOCX_PATH}")

    if not MD_PATH.exists():
        print("ERROR: Markdown file not found."); return

    md_text = MD_PATH.read_text(encoding="utf-8")
    print(f"Read   : {len(md_text):,} characters")

    doc = build_docx(md_text)
    doc.save(str(DOCX_PATH))

    import os
    size_kb = os.path.getsize(DOCX_PATH) / 1024
    print(f"\nDOCX ready : {DOCX_PATH}")
    print(f"Size       : {size_kb:.1f} KB  ({size_kb/1024:.2f} MB)")
    print(f"MD kept    : {MD_PATH.exists()}")


if __name__ == "__main__":
    main()
